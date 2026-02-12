import os
import re
import sys
from pathlib import Path
from typing import List, Tuple, Optional

from dotenv import load_dotenv
import pandas as pd
from docx import Document
from pypdf import PdfReader
from supabase import create_client, Client
from openai import OpenAI


# -----------------------------
# Config
# -----------------------------
EMBED_MODEL = "text-embedding-3-small"
EMBED_DIM = 1536  # for text-embedding-3-small
DEFAULT_CHUNK_CHARS = 400
DEFAULT_OVERLAP_CHARS = 120

SUPPORTED_DOCX = {".docx", ".txt"}
SUPPORTED_PDF = {".pdf"}
SUPPORTED_CSV = {".csv"}


def clean_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # collapse excessive whitespace but keep newlines
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def chunk_text(text: str, chunk_chars: int = DEFAULT_CHUNK_CHARS, overlap_chars: int = DEFAULT_OVERLAP_CHARS) -> List[str]:
    """
    Simple chunker by paragraphs with a sliding overlap.
    Produces chunks ~chunk_chars with overlap to preserve context.
    """
    text = clean_text(text)
    if not text:
        return []

    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[str] = []
    current = ""

    def flush():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for p in paras:
        if len(current) + len(p) + 2 <= chunk_chars:
            current = (current + "\n\n" + p).strip() if current else p
        else:
            flush()
            # if paragraph itself is huge, hard-slice it
            if len(p) > chunk_chars:
                start = 0
                while start < len(p):
                    end = min(start + chunk_chars, len(p))
                    chunks.append(p[start:end].strip())
                    start = max(end - overlap_chars, end)
            else:
                current = p

    flush()

    # Add overlap between chunks (soft overlap)
    if overlap_chars > 0 and len(chunks) > 1:
        overlapped: List[str] = []
        for i, c in enumerate(chunks):
            if i == 0:
                overlapped.append(c)
            else:
                prev = chunks[i - 1]
                overlap = prev[-overlap_chars:] if len(prev) > overlap_chars else prev
                overlapped.append((overlap + "\n\n" + c).strip())
        chunks = overlapped

    # Filter tiny chunks
    chunks = [c for c in chunks if len(c) >= 50]
    return chunks


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Tables -> text rows
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            line = " | ".join([c for c in cells if c])
            if line.strip():
                parts.append(line)

    return clean_text("\n\n".join(parts))

def read_txt(path: Path) -> str:
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: List[str] = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text() or ""
        txt = clean_text(txt)
        if txt:
            parts.append(f"[Page {i+1}]\n{txt}")
    return clean_text("\n\n".join(parts))


def read_csv_as_text(path: Path, max_rows: int = 500) -> str:
    df = pd.read_csv(path)
    if df.empty:
        return ""

    # Limit rows for sanity (you can raise later)
    df = df.head(max_rows)

    # Convert each row to a "fact line" for better retrieval
    lines: List[str] = []
    cols = [str(c) for c in df.columns]
    for _, row in df.iterrows():
        pairs = []
        for c in cols:
            val = row.get(c)
            if pd.isna(val):
                continue
            sval = str(val).strip()
            if sval:
                pairs.append(f"{c}: {sval}")
        if pairs:
            lines.append(" | ".join(pairs))

    title = path.stem.replace("_", " ").strip()
    text = f"{title}\n\n" + "\n".join(lines)
    return clean_text(text)


def make_supabase() -> Client:
    url = os.environ.get("NEXT_PUBLIC_SUPABASE_URL")
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
    if not url:
        raise RuntimeError("Missing env: NEXT_PUBLIC_SUPABASE_URL")
    if not key:
        raise RuntimeError("Missing env: SUPABASE_SERVICE_ROLE_KEY")
    return create_client(url, key)


def make_openai() -> OpenAI:
    key = os.environ.get("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("Missing env: OPENAI_API_KEY")
    return OpenAI(api_key=key)


def insert_document(sb: Client, title: str, source: str) -> str:
    resp = sb.table("documents").insert({"title": title, "source": source}).execute()
    if not resp.data:
        raise RuntimeError(f"Failed to insert document: {title}")
    return resp.data[0]["id"]


def embed_texts(client: OpenAI, texts: List[str]) -> List[List[float]]:
    # OpenAI embeddings API supports batch input
    emb = client.embeddings.create(model=EMBED_MODEL, input=texts)
    vectors = [d.embedding for d in emb.data]
    return vectors


def insert_chunks(sb: Client, document_id: str, chunks: List[str], vectors: List[List[float]]):
    if len(chunks) != len(vectors):
        raise RuntimeError("Chunks and vectors length mismatch")

    rows = []
    for i, (content, vec) in enumerate(zip(chunks, vectors)):
        # basic safety check
        if len(vec) != EMBED_DIM:
            raise RuntimeError(f"Unexpected embedding dim {len(vec)} (expected {EMBED_DIM})")
        rows.append({
            "document_id": document_id,
            "chunk_index": i,
            "content": content,
            "embedding": vec
        })

    # Insert in batches
    batch_size = 100
    for start in range(0, len(rows), batch_size):
        batch = rows[start:start + batch_size]
        resp = sb.table("document_chunks").insert(batch).execute()
        if resp.data is None:
            raise RuntimeError("Failed inserting chunk batch")


def process_file(path: Path) -> Tuple[str, str]:
    """
    Returns: (title, extracted_text)
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        return path.stem, read_docx(path)
    if ext == ".txt":
        return path.stem, read_txt(path)
    if ext in SUPPORTED_PDF:
        return path.stem, read_pdf(path)
    if ext in SUPPORTED_CSV:
        return path.stem, read_csv_as_text(path)
    raise RuntimeError(f"Unsupported file type: {path.name}")


def main():
    # Load .env.local from project root (best-effort)
    # This lets you run without exporting env vars manually.
    root = Path(__file__).resolve().parents[1]
    env_path = root / ".env.local"
    if env_path.exists():
        load_dotenv(env_path)
    else:
        # still allow env vars from shell/Vercel
        load_dotenv()

    sb = make_supabase()
    oa = make_openai()

    sources_dir = root / "rag_sources"
    if not sources_dir.exists():
        raise RuntimeError(f"Missing folder: {sources_dir}")

    inputs: List[Tuple[Path, str]] = []
    for sub, label in [("docs", "docs"), ("pdfs", "pdfs"), ("sheets", "sheets")]:
        d = sources_dir / sub
        if not d.exists():
            continue
        for p in d.rglob("*"):
            if p.is_file():
                ext = p.suffix.lower()
                if ext in SUPPORTED_DOCX or ext in SUPPORTED_PDF or ext in SUPPORTED_CSV:
                    inputs.append((p, label))

    if not inputs:
        print("No files found. Put .docx in rag_sources/docs, .pdf in rag_sources/pdfs, .csv in rag_sources/sheets")
        sys.exit(0)

    print(f"Found {len(inputs)} files to ingest.")

    for path, label in inputs:
        print(f"\n--- Ingesting: {path.name} ({label}) ---")
        title, text = process_file(path)
        if not text or len(text) < 50:
            print("Skipping (no extractable text).")
            continue

        # Chunk
        chunks = chunk_text(text, DEFAULT_CHUNK_CHARS, DEFAULT_OVERLAP_CHARS)
        if not chunks:
            print("Skipping (no chunks produced).")
            continue

        # Insert doc
        doc_id = insert_document(sb, title=title, source=f"{label}:{path.name}")
        print(f"Document inserted: {doc_id} | chunks={len(chunks)}")

        # Embed + insert chunks
        vectors = embed_texts(oa, chunks)
        insert_chunks(sb, document_id=doc_id, chunks=chunks, vectors=vectors)
        print("Chunks embedded + stored.")

    print("\nDone.")


if __name__ == "__main__":
    main()
